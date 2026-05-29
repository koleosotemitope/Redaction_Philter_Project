import argparse
import re
from html import unescape
from pathlib import Path
from typing import List, Tuple


def read_docx_text(path: Path) -> Tuple[str, str]:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document  # type: ignore
    except Exception:
        return "", "missing python-docx"

    try:
        doc = Document(str(path))
        chunks: List[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                chunks.append(para.text)
        text = "\n".join(chunks).strip()
        return text, "ok" if text else "empty"
    except Exception as exc:
        return "", f"docx parse error: {exc}"


def read_doc_text(path: Path) -> Tuple[str, str]:
    """Extract text from a DOC file using python-docx (newer .doc files) or fallback to OCR."""
    try:
        from docx import Document  # type: ignore
    except Exception:
        return "", "missing python-docx"

    try:
        doc = Document(str(path))
        chunks: List[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                chunks.append(para.text)
        text = "\n".join(chunks).strip()
        return text, "ok" if text else "empty"
    except Exception as exc:
        return "", f"doc parse error: {exc}"


def read_html_text(path: Path) -> Tuple[str, str]:
    """Extract readable text from an HTML/HTM file without extra dependencies."""
    try:
        raw = path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        return "", f"html read error: {exc}"

    try:
        cleaned = re.sub(r"(?is)<(script|style)\b.*?>.*?</\1>", " ", raw)
        cleaned = re.sub(r"(?is)<br\s*/?>", "\n", cleaned)
        cleaned = re.sub(r"(?is)</p\s*>", "\n", cleaned)
        cleaned = re.sub(r"(?is)<[^>]+>", " ", cleaned)
        cleaned = unescape(cleaned)
        cleaned = re.sub(r"\r\n?", "\n", cleaned)
        cleaned = re.sub(r"[ \t]+", " ", cleaned)
        cleaned = re.sub(r"\n\s+", "\n", cleaned)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        text = cleaned.strip()
        return text, "ok" if text else "empty"
    except Exception as exc:
        return "", f"html parse error: {exc}"


def read_pdf_text(path: Path) -> Tuple[str, str]:
    """Extract text from a PDF via pypdf; fallback OCR is handled separately."""
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception:
        return "", "missing pypdf"

    try:
        reader = PdfReader(str(path))
        chunks: List[str] = []
        for page in reader.pages:
            chunks.append(page.extract_text() or "")
        text = "\n".join(chunks).strip()
        return text, "ok" if text else "empty"
    except Exception as exc:
        return "", f"pdf parse error: {exc}"


def ocr_pdf(path: Path) -> Tuple[str, str]:
    """OCR a PDF by rasterizing pages, requires pdf2image + pytesseract + tesseract binary."""
    try:
        import pytesseract  # type: ignore
        from pdf2image import convert_from_path  # type: ignore
    except Exception:
        return "", "missing pdf2image/pytesseract"

    try:
        images = convert_from_path(str(path))
        chunks: List[str] = []
        for img in images:
            chunks.append(pytesseract.image_to_string(img))
        text = "\n".join(chunks).strip()
        return text, "ok" if text else "empty"
    except Exception as exc:
        return "", f"pdf ocr error: {exc}"


def ocr_image(path: Path) -> Tuple[str, str]:
    """OCR an image file, requires pillow + pytesseract + tesseract binary."""
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except Exception:
        return "", "missing pillow/pytesseract"

    try:
        with Image.open(path) as img:
            text = pytesseract.image_to_string(img).strip()
        return text, "ok" if text else "empty"
    except Exception as exc:
        return "", f"image ocr error: {exc}"


def collect_files(input_path: Path, recursive: bool) -> List[Path]:
    exts = {".pdf", ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".html", ".htm"}
    if input_path.is_file():
        return [input_path] if input_path.suffix.lower() in exts else []

    globber = input_path.rglob if recursive else input_path.glob
    return sorted([p for p in globber("*") if p.is_file() and p.suffix.lower() in exts])


def convert_one(path: Path, out_dir: Path) -> Tuple[bool, str, Path]:
    suffix = path.suffix.lower()
    text = ""
    status = ""

    if suffix == ".pdf":
        text, status = read_pdf_text(path)
        if not text:
            text, status = ocr_pdf(path)
    elif suffix in {".html", ".htm"}:
        text, status = read_html_text(path)
    else:
        text, status = ocr_image(path)

    output_file = out_dir / f"{path.stem}.txt"
    if text:
        output_file.write_text(text + "\n", encoding="utf-8")
        return True, status, output_file

    return False, status, output_file


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert PDF, HTML, and image files into plain text files for Philter input."
    )
    parser.add_argument(
        "-i",
        "--input",
        required=True,
        help="Input file or directory containing PDF/HTML/JPEG/PNG/etc.",
    )
    parser.add_argument(
        "-o",
        "--output",
        default="./data/ingested_txt",
        help="Output directory for generated .txt files.",
    )
    parser.add_argument(
        "--no-recursive",
        action="store_true",
        help="Do not recurse through subdirectories when input is a folder.",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    input_path = Path(args.input)
    out_dir = Path(args.output)

    if not input_path.exists():
        print(f"Input path not found: {input_path}")
        return 1

    out_dir.mkdir(parents=True, exist_ok=True)

    files = collect_files(input_path, recursive=not args.no_recursive)
    if not files:
        print("No supported files found. Supported: .pdf, .html, .htm, .png, .jpg, .jpeg, .tif, .tiff, .bmp")
        return 1

    ok_count = 0
    fail_count = 0

    for file_path in files:
        ok, status, out_file = convert_one(file_path, out_dir)
        if ok:
            ok_count += 1
            print(f"OK   {file_path} -> {out_file} ({status})")
        else:
            fail_count += 1
            print(f"FAIL {file_path} ({status})")

    print(f"Done. Converted: {ok_count}, Failed: {fail_count}, Output dir: {out_dir}")

    if fail_count > 0:
        print("Tip: For OCR support install pillow, pytesseract, pdf2image and system Tesseract OCR.")

    return 0 if ok_count > 0 else 1


if __name__ == "__main__":
    raise SystemExit(main())
